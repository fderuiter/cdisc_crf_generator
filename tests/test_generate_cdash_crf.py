import subprocess
import sys


def test_help():
    result = subprocess.run(
        [sys.executable, "scripts/generate_cdash_crf.py", "--help"],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0, result.stderr
    assert "Generate Word CRFs from CDASH metadata workbooks." in result.stdout


def test_generate(tmp_path):
    out_dir = tmp_path / "out"
    result = subprocess.run(
        [
            sys.executable,
            "scripts/generate_cdash_crf.py",
            "--model",
            "tests/CDASH_Model_v1.3.xlsx",
            "--ig",
            "tests/CDASHIG_v2.3 (1).xlsx",
            "--out",
            str(out_dir),
            "--domains",
            "AE",
        ],
        capture_output=True,
        text=True,
    )
    assert result.returncode == 0, result.stderr
    doc_path = out_dir / "AE_CRF.docx"
    assert doc_path.exists()

    from docx import Document
    from zipfile import ZipFile

    doc = Document(doc_path)
    table = doc.tables[0]
    assert len(table.columns) == 7
    assert table.cell(0, 6).text == "Required"

    texts = "\n".join(p.text for p in doc.paragraphs)
    assert "Footnotes" in texts
    assert "[1]" in texts

    admin = doc.tables[1].cell(0, 0).text.replace("\xa0", " ")
    assert admin == "SECTION A  ADMINISTRATIVE"

    with ZipFile(doc_path) as zf:
        xml = zf.read("word/document.xml").decode("utf-8")
        assert "w14:checkbox" in xml or "w14:date" in xml
        assert "Validate dependencies" in xml
        assert xml.count("<w:bottom") > 0
