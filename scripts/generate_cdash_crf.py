#!/usr/bin/env python
"""
Generate MS‑Word CRF shells from CDASH metadata workbooks
-------------------------------------------------------
This **rewrite** produces a layout that more closely resembles the dark, sectioned
mock‑up supplied in the screenshot:

*   Landscape orientation, dark background header/footer bands, white text.
*   A two‑row page header with room for study‑level information plus the CRF
    title (Domain description written‑out as *Assessment*, *Concomitant / Prior
    Medications*, etc.).
*   Automatic page numbering in the footer (right aligned) and a version label
    (left aligned) just like the mock‑up.
*   Distinct colour‑banded *SECTION A ADMINISTRATIVE* & *SECTION B <DOMAIN>*
    table headers.
*   Six data columns – Variable | Label / Question | Type | Controlled
    Terminology | Data Entry | Instructions – identical to the original script
    but with shading and stronger typographic hierarchy.
*   A helper routine to shade table cells (python‑docx still lacks a high‑level
    API for this).
*   A consolidated mapping of CDASH domain codes → “[category, full title]” so
    we can replace cryptic two‑letter codes with human‑friendly names wherever
    appropriate.

The goal is *not* to pixel‑match the example (python‑docx cannot create shapes
or literal check‑boxes) but to deliver a genuinely usable CRF template that is
visually much closer to the supplied design while remaining fully generated
from the metadata.
"""

import argparse
import pathlib
from typing import Dict, Tuple

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt, RGBColor

###############################################################################
# Domain‑to‑category mapping
###############################################################################

DOMAIN_INFO: Dict[str, Tuple[str, str]] = {
    # Interventions
    "AG": ("Interventions", "Procedure Agents"),
    "CM": ("Interventions", "Concomitant / Prior Medications"),
    "EC": ("Interventions", "Exposure as Collected"),
    "EX": ("Interventions", "Exposure"),
    "ML": ("Interventions", "Meal Data"),
    "PR": ("Interventions", "Procedures"),
    "SU": ("Interventions", "Substance Use"),

    # Events
    "AE": ("Events", "Adverse Events"),
    "CE": ("Events", "Clinical Events"),
    "DS": ("Events", "Disposition"),
    "DV": ("Events", "Protocol Deviations"),
    "HO": ("Events", "Healthcare Encounters"),
    "MH": ("Events", "Medical History"),
    "SA": ("Events", "Serious Adverse Events"),

    # Findings
    "CP": ("Findings", "Cell Phenotype Findings"),
    "CV": ("Findings", "Cardiovascular System Findings"),
    "DA": ("Findings", "Product Accountability"),
    "DD": ("Findings", "Death Details"),
    "ED": ("Findings", "Central Reading"),
    "GF": ("Findings", "Genomics Findings"),
    "IE": ("Findings", "Inclusion / Exclusion Criteria Not Met"),
    "LB": ("Findings", "Laboratory Test Results"),
    "MB": ("Findings", "Microbiology Specimen"),
    "MI": ("Findings", "Microscopic Findings"),
    "MK": ("Findings", "Musculoskeletal System Findings"),
    "MS": ("Findings", "Microbiology Susceptibility"),
    "NV": ("Findings", "Nervous System Findings"),
    "OE": ("Findings", "Ophthalmic Examinations"),
    "PC": ("Findings", "Pharmacokinetics Concentrations"),
    "PE": ("Findings", "Physical Examination"),
    "RE": ("Findings", "Respiratory System Findings"),
    "RP": ("Findings", "Reproductive System Findings"),
    "RS": ("Findings", "Disease Response & Clinical Classification"),
    "SC": ("Findings", "Subject Characteristics"),
    "TR": ("Findings", "Tumor / Lesion Results"),
    "TU": ("Findings", "Tumor / Lesion Identification"),
    "UR": ("Findings", "Urinary System Findings"),
    "VS": ("Findings", "Vital Signs"),

    # Findings‑About
    "FA": ("Findings About", "Findings About Events or Interventions"),
    "SR": ("Findings About", "Skin Response"),

    # Special Purpose
    "CO": ("Special Purpose", "Comments"),
    "DM": ("Special Purpose", "Demographics"),
}

def get_domain_info(domain: str) -> Tuple[str, str]:
    """Return (category, full title) for *domain* code."""
    try:
        return DOMAIN_INFO[domain.upper()]
    except KeyError:
        return ("Unknown", domain)

###############################################################################
# Low‑level helpers
###############################################################################

def _add_page_field(paragraph):
    """Insert Word PAGE field into *paragraph* (in‑place)."""
    run = paragraph.add_run()

    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    run._r.append(instr)

    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)


def _set_cell_shading(cell, color_hex: str = "4F81BD"):
    """Shade *cell* background with *color_hex* (RGB string without #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    # Remove existing shading if any
    for shd in tc_pr.findall("w:shd", tc_pr.nsmap):
        tc_pr.remove(shd)
    # Add new shading element
    shd_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    tc_pr.append(shd_elm)


def _style_header_cell(cell):
    """Apply white bold text to header *cell*."""
    para = cell.paragraphs[0]
    run = para.runs[0]
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

###############################################################################
# Data I/O helpers
###############################################################################

def load_ig(ig_path: str) -> pd.DataFrame:
    """Load and normalise the *Variables* worksheet from a CDASH IG workbook."""
    ig_df = pd.read_excel(ig_path, sheet_name="Variables", engine="openpyxl")
    ig_df = ig_df[~ig_df["Domain"].isna()].copy()

    ig_df["Display Label"] = ig_df["Question Text"].fillna(
        ig_df["CDASHIG Variable Label"]
    )

    ig_df.rename(
        columns={
            "CDASHIG Variable": "Variable",
            "Variable Order": "Order",
            "Case Report Form Completion Instructions": "CRF Instructions",
            "CDISC CT Codelist Submission Values(s), Subset Submission Value(s)": "CT Values",
            "CDISC CT Codelist Code(s), Subset Codes(s)": "CT Codes",
        },
        inplace=True,
    )

    return ig_df

###############################################################################
# Core CRF builder
###############################################################################

def build_domain_crf(domain_df: pd.DataFrame, domain: str, out_dir: pathlib.Path) -> None:
    """Build a Word document for a single CDASH *domain* and save it to disk."""

    category, full_title = get_domain_info(domain)

    # ---------------------------------------------------------------------
    #  Document meta & base formatting
    # ---------------------------------------------------------------------
    document = Document()

    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    # Uniform font for entire document
    style = document.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ---------------------------------------------------------------------
    #  Page header (dark band with white text)
    # ---------------------------------------------------------------------
    header = section.header

    hdr_tbl = header.add_table(rows=2, cols=2, width=section.page_width)
    hdr_tbl.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_tbl.autofit = False
    hdr_tbl.repeat_rows = 0

    sponsor_cell, title_cell = hdr_tbl.rows[0].cells
    meta_cell_L, meta_cell_R = hdr_tbl.rows[1].cells

    # Row‑0: sponsor block & CRF title block
    sponsor_cell.text = "Sponsor Study Name"
    title_cell.text = full_title

    sponsor_cell.width = title_cell.width = section.page_width / 2

    for c in (sponsor_cell, title_cell):
        _set_cell_shading(c, "1F1F1F")  # very dark grey
        _style_header_cell(c)

    title_para = title_cell.paragraphs[0]
    title_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    title_para.runs[0].font.size = Pt(14)

    # Row‑1: Subject meta‑data placeholders
    meta_cell_L.text = "Subject ID: _____‑___‑___    SITE #: ___"
    meta_cell_R.text = "Initials: ___ ___ ___"
    for c in (meta_cell_L, meta_cell_R):
        _set_cell_shading(c, "3F3F3F")
        _style_header_cell(c)

    # ---------------------------------------------------------------------
    #  Footer (version label at left, page # at right)
    # ---------------------------------------------------------------------
    footer = section.footer

    f_left = footer.add_paragraph(f"{full_title}, Version 1.0 DRAFT")
    f_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    f_left.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    f_right = footer.add_paragraph()
    f_right.paragraph_format.right_indent = Pt(0)
    _add_page_field(f_right)
    f_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    f_right.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # ---------------------------------------------------------------------
    #  SECTION A – ADMINISTRATIVE (static content)
    # ---------------------------------------------------------------------
    document.add_paragraph()
    secA_tbl = document.add_table(rows=3, cols=2, style="Table Grid")
    secA_tbl.autofit = False
    secA_tbl.allow_autofit = False

    # Header row (spanning two columns)
    hdr_row = secA_tbl.rows[0]
    hdr_cell = hdr_row.cells[0]
    hdr_cell.merge(hdr_row.cells[1])
    hdr_cell.text = "SECTION A    ADMINISTRATIVE"
    _set_cell_shading(hdr_cell, "8064A2")  # muted purple
    _style_header_cell(hdr_cell)

    # Row 1 – Question completed?
    secA_tbl.rows[1].cells[0].text = f"Was {full_title.lower()} completed?"
    secA_tbl.rows[1].cells[1].text = "○ No (Complete protocol deviation form)    ○ Yes"

    # Row 2 – Date of assessment
    secA_tbl.rows[2].cells[0].text = "Date of assessment:"
    secA_tbl.rows[2].cells[1].text = "__|__|____|____|    DD‑MMM‑YYYY"

    # ---------------------------------------------------------------------
    #  SECTION B – DOMAIN VARIABLES
    # ---------------------------------------------------------------------
    document.add_paragraph()
    var_tbl = document.add_table(rows=1, cols=6, style="Table Grid")
    var_tbl.autofit = False

    hdr_cells = var_tbl.rows[0].cells
    col_titles = [
        "Variable",
        "Label / Question",
        "Type",
        "Controlled Terminology",
        "Data Entry",
        "Instructions",
    ]
    for idx, title in enumerate(col_titles):
        hdr_cells[idx].text = title
        _set_cell_shading(hdr_cells[idx], "4F81BD")
        _style_header_cell(hdr_cells[idx])

    # Data rows ordered by the "Variable Order" column
    for _, row in domain_df.sort_values("Order").iterrows():
        cells = var_tbl.add_row().cells
        # 0 Variable name
        cells[0].text = row["Variable"]

        # 1 Label / Question
        cells[1].text = str(row["Display Label"])

        # 2 Data type (where available)
        cells[2].text = str(row.get("Type", ""))

        # 3 Controlled terminology – prefer values over codes
        ct_val = row.get("CT Values")
        ct_code = row.get("CT Codes")
        cells[3].text = (
            str(ct_val)
            if pd.notna(ct_val)
            else str(ct_code) if pd.notna(ct_code) else ""
        )

        # 4 Data entry placeholder (simple underline for now)
        cells[4].text = "_______________"

        # 5 Instructions (italic, stacked if multiple)
        instructions = []
        if pd.notna(row.get("CRF Instructions")):
            instructions.append(str(row.get("CRF Instructions")))
        if pd.notna(row.get("Implementation Notes")):
            instructions.append(str(row.get("Implementation Notes")))

        # Auto‑detect date fields and add a formatting hint
        label_lower = str(row.get("Display Label", "")).lower()
        var_upper = row["Variable"].upper()
        if "date" in label_lower or var_upper.endswith(("DT", "DAT")):
            instructions.append("Format: dd/mm/yyyy")

        instr_para = cells[5].paragraphs[0]
        for idx, item in enumerate(instructions):
            run = instr_para.add_run(item)
            run.italic = True
            if idx < len(instructions) - 1:
                instr_para.add_run("\n")

    # ---------------------------------------------------------------------
    #  Save document
    # ---------------------------------------------------------------------
    out_path = out_dir / f"{domain}_{full_title.replace(' ', '_')}_CRF.docx"
    document.save(out_path)
    print(f"\u2713 Saved {out_path.relative_to(out_dir.parent)}")

###############################################################################
# CLI entry‑point
###############################################################################

def main() -> None:
    parser = argparse.ArgumentParser(
        description="Generate Word CRF shells from CDASH metadata workbooks."
    )
    parser.add_argument("--model", required=True, help="Path to CDASH_Model_v1.3.xlsx (reserved for future use)")
    parser.add_argument("--ig", required=True, help="Path to CDASHIG_v2.3.xlsx")
    parser.add_argument("--out", default="crfs", help="Directory for generated Word documents")
    parser.add_argument(
        "--domains", nargs="*", metavar="DOMAIN", help="Optional domain whitelist (e.g. AE CM VS)"
    )

    args = parser.parse_args()
    out_dir = pathlib.Path(args.out)
    out_dir.mkdir(parents=True, exist_ok=True)

    ig_df = load_ig(args.ig)

    target_domains = [d.upper() for d in (args.domains or ig_df["Domain"].unique())]
    for dom in target_domains:
        dom_df = ig_df[ig_df["Domain"] == dom]
        if dom_df.empty:
            print(f"\u26A0 Domain {dom} not found in IG – skipped")
            continue
        build_domain_crf(dom_df, dom, out_dir)


if __name__ == "__main__":
    main()
